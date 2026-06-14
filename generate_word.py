import docx
from docx.shared import Pt
import re

def create_word():
    doc = docx.Document()
    
    # Read the markdown
    with open("Laporan_Akademis.md", "r", encoding="utf-8") as f:
        content = f.read()
        
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", "").replace("*", ""), 0)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", "").replace("*", ""), 1)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", "").replace("*", ""), 2)
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", "").replace("*", ""), style='List Bullet')
        elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. "):
            doc.add_paragraph(line[3:].replace("*", ""), style='List Number')
        elif line == "---":
            continue
        elif line.startswith("$$") or line.startswith("$"):
            # Math formulas, just add as text
            p = doc.add_paragraph()
            p.add_run(line.replace("$", "")).italic = True
        else:
            # Normal paragraph
            # Basic bold removal for plain text
            text = line.replace("**", "").replace("*", "")
            doc.add_paragraph(text)
            
    doc.save("Laporan_Akademis.docx")
    print("✅ Berhasil! File 'Laporan_Akademis.docx' sudah dibuat di folder proyek Anda.")

if __name__ == "__main__":
    create_word()
